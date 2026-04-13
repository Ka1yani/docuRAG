import os
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy.orm import Session
from sqlalchemy import text
from app.models import Document, DocumentChunk
from app.services.llm_service import get_embedding
from app.logger import logger

try:
    from faster_whisper import WhisperModel
except ImportError:
    WhisperModel = None

try:
    import easyocr
except ImportError:
    easyocr = None

def format_timestamp(seconds: float) -> str:
    mins = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{mins:02d}:{secs:02d}"

def chunk_text(text_content: str, word_chunk_size: int = 400) -> list[str]:
    words = text_content.split()
    chunks = []
    for i in range(0, len(words), word_chunk_size):
        chunk = " ".join(words[i : i + word_chunk_size])
        if chunk.strip():
            chunks.append(chunk)
    return chunks

def process_and_store_document(file_path: str, file_name: str, db: Session):
    # 1. Create Document record
    db_doc = Document(file_name=file_name)
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)

    _, ext = os.path.splitext(file_name.lower())
    logger.info(f"Extracting {ext} logic for {file_name}")
    
    extracted_data = [] # list of (page_num, text)

    # 2. Extract Text
    if ext == ".pdf":
        reader = PdfReader(file_path)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text_content = page.extract_text()
            if text_content and text_content.strip():
                extracted_data.append((page_num + 1, text_content))
    elif ext in [".doc", ".docx"]:
        doc = DocxDocument(file_path)
        # Using paragraph index as "page" since word wrapping is dynamic
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                extracted_data.append((i + 1, para.text))
    elif ext in [".mp3", ".wav", ".m4a"]:
        if WhisperModel is None:
            raise ImportError("faster-whisper is not installed. Unable to process audio.")
        
        logger.info(f"Initializing Whisper for audio transcription: {file_name}")
        # Use CPU by default. Set device="cuda" if running on a GPU.
        # compute_type="int8" reduces memory usage for CPU.
        model = WhisperModel("base", device="cpu", compute_type="int8")
        
        segments, info = model.transcribe(file_path, beam_size=5)
        logger.info(f"Detected language '{info.language}' with probability {info.language_probability}")
        
        for i, segment in enumerate(segments):
            start_str = format_timestamp(segment.start)
            end_str = format_timestamp(segment.end)
            text_val = segment.text.strip()
            
            # Format text explicitly to bias the LLM with audio flow context
            formatted_text = f"[Audio Segment {start_str} - {end_str}]: {text_val}"
            
            # Use segment ordinal as the logical page number
            extracted_data.append((i + 1, formatted_text))
            
    elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
        if easyocr is None:
            raise ImportError("easyocr is not installed. Unable to process image.")
            
        logger.info(f"Initializing EasyOCR for image extraction: {file_name}")
        # Initialize reader natively on CPU for universal compatibility. 
        # EasyOCR will automatically use GPU if PyTorch CUDA is bound.
        reader = easyocr.Reader(['en'], gpu=False)
        
        logger.info("Reading raw text from image structural patterns...")
        results = reader.readtext(file_path, detail=0)
        
        # Merge all detected text bounding boxes into a cohesive string
        extracted_text = "\n".join(results)
        
        if extracted_text.strip():
            # Apply strict semantic framing to bias the LLM contextual window
            formatted_text = f"[Transcribed Image Content]:\n{extracted_text}"
            extracted_data.append((1, formatted_text))
        else:
            logger.info("No text detected in the image.")

    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text_content = f.read()
            extracted_data.append((1, text_content))
    else:
        logger.error(f"Unsupported extraction attempted: {ext}")
        raise ValueError(f"Unsupported file extension: {ext}")

    # 3. Chunk and Store
    for page_num, text_content in extracted_data:
        # Clean text
        clean_text = " ".join(text_content.replace("\\n", " ").split())
        if not clean_text:
            continue
            
        chunks = chunk_text(clean_text)
        for chunk in chunks:
            embedding = get_embedding(chunk)
            db_chunk = DocumentChunk(
                document_id=db_doc.id,
                file_name=file_name,
                page_number=page_num,
                content=chunk,
                content_vector=str(embedding) if embedding else None
            )
            db.add(db_chunk)
            
    db.commit()
    logger.info(f"Completed! {file_name} split into database chunks with vector embeddings attached.")
    return db_doc
